from docx import Document

def crear_doc():
    doc = Document()
    doc.add_heading("Prueba", 0)
    doc.add_paragraph("¡Funciona en Railway!")
    doc.save("prueba.docx")

crear_doc()
